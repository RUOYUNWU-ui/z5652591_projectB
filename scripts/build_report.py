"""Build an evidence-driven report draft without overwriting the final report.

The submitted ``report/report.docx`` contains the student's final manual edits.
This script deliberately writes a separate reproducible draft so rerunning it
cannot destroy that authored source.
"""
from __future__ import annotations

import pathlib

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = pathlib.Path(__file__).resolve().parent.parent
REPORT_DIR = ROOT / "report"
FIGURE_DIR = ROOT / "results" / "figures"
TABLE_DIR = ROOT / "results" / "tables"
OUTPUT = REPORT_DIR / "report_generated_draft.docx"

NAVY = "2C3E50"
BLUE = "315B7D"
GOLD = "B98A3D"
MUTED = "5B6570"
PALE = "F4F6F9"
WHITE = "FFFFFF"


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_fixed_table(table, widths: list[float]) -> None:
    """Set fixed Word-native table geometry; widths must total 6.5 inches."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = value
        shade(cell, NAVY)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor.from_string(WHITE)
            run.bold = True
            run.font.size = Pt(8)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
            if row_index % 2:
                shade(cells[index], PALE)
            for run in cells[index].paragraphs[0].runs:
                run.font.size = Pt(8)
    set_fixed_table(table, widths)
    return table


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(MUTED)


def add_figure(doc: Document, filename: str, caption: str, width: float = 6.35) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(FIGURE_DIR / filename), width=Inches(width))
    add_caption(doc, caption)


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.22
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, BLUE),
        ("Heading 2", 13, 12, 6, BLUE),
        ("Heading 3", 12, 8, 4, NAVY),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("SIGNALYIELD  |  FINS5545 PART B")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Page ")
    add_field(footer, "PAGE")
    footer.add_run(" of ")
    add_field(footer, "NUMPAGES")
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(MUTED)


def add_cover(doc: Document) -> None:
    doc.add_paragraph().paragraph_format.space_after = Pt(92)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("FINS5545  |  FINTECH PROJECT 2026")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(GOLD)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("SignalYield")
    run.bold = True
    run.font.size = Pt(31)
    run.font.color.rgb = RGBColor.from_string(NAVY)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(30)
    run = subtitle.add_run(
        "Systematic Multi-Asset Funds with News-Sentiment Analytics"
    )
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    for text, bold in (
        ("Part B report", True),
        ("Student: z5652591", False),
        ("Sample: 2020-01-01 to 2023-12-31 | OOS: 2021-2023", False),
        ("Generated evidence build: 14 August 2026", False),
    ):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(text)
        run.font.size = Pt(10.5)
        run.bold = bold
        run.font.color.rgb = RGBColor.from_string(MUTED)

    doc.add_paragraph().paragraph_format.space_after = Pt(48)
    note = doc.add_table(rows=1, cols=1)
    note.style = "Table Grid"
    note.cell(0, 0).text = (
        "AI-assisted editable draft. The 50-headline and 24-term reviews are complete; "
        "the student must still verify every number and rewrite the final economic "
        "interpretation in their own words before submission."
    )
    shade(note.cell(0, 0), "FFF5DE")
    set_fixed_table(note, [6.5])
    for run in note.cell(0, 0).paragraphs[0].runs:
        run.font.size = Pt(9)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(NAVY)
    doc.add_page_break()


def pct(value: float) -> str:
    return f"{value:.1%}"


def build() -> pathlib.Path:
    metrics = pd.read_csv(TABLE_DIR / "performance_metrics.csv")
    fusion = pd.read_csv(TABLE_DIR / "fusion_comparison.csv")
    sentiment = pd.read_csv(TABLE_DIR / "sentiment_model_comparison.csv")
    shrinkage = pd.read_csv(TABLE_DIR / "shrinkage_comparison.csv")
    costs = pd.read_csv(TABLE_DIR / "transaction_cost_sensitivity.csv")
    fund_weights = pd.read_csv(TABLE_DIR.parent / "data" / "fund_weights.csv")
    sector_index = pd.read_csv(TABLE_DIR.parent / "data" / "sector_sentiment_index.csv")
    manual_validation = pd.read_csv(
        TABLE_DIR / "sentiment_manual_review_validation.csv"
    )
    manual_review = pd.read_csv(TABLE_DIR / "sentiment_manual_review_template.csv")
    lexicon_review = pd.read_csv(TABLE_DIR / "finance_lexicon_candidates.csv")

    metric_by_fund = metrics.set_index("fund_id")
    min_var = metrics.loc[metrics["method"].eq("min_variance")].set_index(
        "asset_family"
    )
    combined_crypto = fund_weights.loc[
        fund_weights["asset_class"].eq("crypto")
        & fund_weights["fund_id"].str.startswith("combined_")
    ]
    mean_crypto_weight = combined_crypto.groupby("fund_id")["weight"].sum().div(
        combined_crypto["effective_date"].nunique()
    )
    enhanced_sector = sector_index.loc[
        sector_index["model"].eq("vader_finance_enhanced")
    ]
    sector_summary = enhanced_sector.groupby("sector").agg(
        mean_sentiment=("sentiment", "mean"),
        mean_headlines=("n_headlines", "mean"),
    )

    doc = Document()
    configure(doc)
    add_cover(doc)

    doc.add_heading("Executive summary", level=1)
    doc.add_paragraph(
        "SignalYield is an educational investment-product prototype for a financially "
        "literate user who wants transparent, comparable access to systematic equity, "
        "crypto and multi-asset funds. The product converts the audited Part A data "
        "foundation into 12 investable fund simulations, a sector news-sentiment index, "
        "and an investor dashboard. Every fund is evaluated using a monthly expanding-"
        "window walk-forward design: information available by a rebalance date sets "
        "weights that become effective on the next trading day, after which holdings "
        "drift with realised returns until the next monthly trade."
    )
    best = metrics.loc[metrics["sharpe"].idxmax()]
    doc.add_paragraph(
        f"The strongest baseline risk-adjusted result was {best['fund_id']} with an "
        f"annualised Sharpe ratio of {best['sharpe']:.2f}, annualised return of "
        f"{pct(best['annualised_return'])}, and maximum drawdown of "
        f"{pct(best['max_drawdown'])}. Crypto funds delivered high returns but also "
        "exceptionally high volatility and drawdowns, so the evidence does not support "
        "presenting crypto as a low-risk substitute for equities. Combined minimum "
        "variance was almost identical to equity minimum variance because the optimiser "
        "assigned little economic value to volatile crypto exposure."
    )
    doc.add_paragraph(
        "The finance-enhanced VADER model reduced the neutral-headline rate from "
        f"{sentiment.loc[sentiment['model'].eq('vader_plain'), 'neutral_rate'].iat[0]:.1%} "
        f"to {sentiment.loc[sentiment['model'].eq('vader_finance_enhanced'), 'neutral_rate'].iat[0]:.1%}. "
        "A fixed, untuned sentiment fusion improved the minimum-variance equity fund in "
        "this sample, especially for the contrarian direction, but this is exploratory "
        "rather than proof of a stable forecasting relation. Two robustness extensions "
        "add practical depth: Ledoit-Wolf covariance shrinkage and a 0-100 basis-point "
        "transaction-cost curve. A completed 50-headline human review improved overall "
        "agreement from 66% for plain VADER to 70% for the enhanced model."
    )
    doc.add_heading("1. Product and investor proposition", level=1)
    doc.add_paragraph(
        "SignalYield addresses a presentation gap rather than claiming a new source of "
        "guaranteed alpha. A user can compare consistent OOS metrics, inspect a fund's "
        "latest target holdings, combine precomputed funds in an allocation builder, "
        "and view sector news sentiment as context. The interface avoids raw-data and "
        "model execution at runtime: it reads committed CSV evidence, making the app "
        "faster, cheaper to host, and easier to audit. The commercial analogy is a "
        "systematic fund platform earning a management fee; in this project it remains "
        "an educational prototype and not personal financial advice."
    )
    doc.add_paragraph(
        "The product design deliberately separates evidence from choice. Performance "
        "tables describe historical simulations; the fact sheet explains one fund; the "
        "allocation builder performs only a linear combination of saved OOS returns; "
        "and sentiment is labelled as noisy context. This reduces the risk that visual "
        "polish is mistaken for predictive certainty."
    )

    doc.add_heading("2. Funds and backtest design", level=1)
    doc.add_paragraph(
        "The fund universe is a three-by-four matrix: equity-only, crypto-only and "
        "combined equity-plus-crypto universes crossed with equal weight, minimum "
        "variance, maximum Sharpe and risk parity. All optimised funds are long-only "
        "and fully invested. Equal weight is an important estimation-free benchmark; "
        "minimum variance uses the covariance matrix; maximum Sharpe uses estimated "
        "mean returns and covariance with a zero risk-free rate; and risk parity seeks "
        "equal contributions to portfolio variance. These choices follow the mean-"
        "variance diversification tradition [1] while retaining a simple benchmark "
        "that can be hard to beat after estimation error [3]."
    )
    doc.add_paragraph(
        "Equity and combined funds require 252 trading-day observations before their "
        "first live return; crypto requires 365 calendar days. Estimation expands from "
        "the start of the sample, rebalancing occurs monthly, and each new weight is "
        "effective one trading day after estimation. The first live dates are 5 January "
        "2021 for equity and combined funds and 2 January 2021 for crypto. An automated "
        "audit confirms estimation_end is strictly earlier than effective_date for all "
        "432 fund-rebalance records. Equity and combined statistics use 252-day "
        "annualisation; crypto uses 365 days. Baseline transaction costs are zero. "
        "Holdings drift between monthly effective dates; turnover is half the absolute "
        "difference between drifted pre-trade and new target weights."
    )
    add_table(
        doc,
        ["Design choice", "Implementation", "Reason"],
        [
            ["Window", "Expanding", "Uses growing history without discarding early data"],
            ["Rebalance", "Monthly", "Balances responsiveness and implementation burden"],
            ["Timing", "Weights live next day", "Prevents same-day information leakage"],
            ["Constraints", "Long-only; sum to 1", "Clear and investable educational mandate"],
            ["Risk-free rate", "0", "Transparent Sharpe convention allowed by brief"],
        ],
        [1.25, 1.65, 3.60],
    )
    add_caption(doc, "Table 1. Frozen baseline backtest design, specified before OOS evaluation.")
    doc.add_page_break()

    doc.add_heading("3. Out-of-sample fund results", level=1)
    doc.add_paragraph(
        "Figure A1 and Table A1 report all 12 funds. The growth paths show that the "
        "fastest-growing crypto funds also experienced the widest fluctuations, while "
        "equity and combined minimum variance followed the smoothest paths. Equity equal "
        f"weight achieved a Sharpe ratio of {metric_by_fund.loc['equity_equal_weight','sharpe']:.2f}, "
        "ahead of the three optimised equity alternatives. This "
        "is economically plausible: estimating means and covariances from a short, "
        "non-stationary sample can introduce enough error to offset theoretical gains. "
        "Maximum Sharpe was the weakest equity method on risk-adjusted performance and "
        "also had the highest turnover, consistent with sensitivity to estimated means."
    )
    doc.add_paragraph(
        f"Crypto risk parity and equal weight produced annualised returns of "
        f"{pct(metric_by_fund.loc['crypto_risk_parity','annualised_return'])} and "
        f"{pct(metric_by_fund.loc['crypto_equal_weight','annualised_return'])}, but their "
        "annualised volatility was about 80%-82% and maximum drawdowns exceeded 80%. "
        "Figure A2 isolates minimum variance: crypto still lost "
        f"{abs(min_var.loc['crypto','max_drawdown']):.1%} from peak to trough versus "
        f"{abs(min_var.loc['equity','max_drawdown']):.1%} for equity. The equity and "
        "combined lines overlap because their minimum-variance allocations are nearly "
        "identical. These results make a critical product distinction: "
        "high ending wealth does not imply a tolerable investor journey. The dashboard "
        "therefore shows drawdown and volatility alongside return."
    )
    doc.add_paragraph(
        "Figure A3 explains the combined-fund allocation mechanism. Average crypto "
        f"exposure was {mean_crypto_weight.loc['combined_equal_weight']:.1%} for equal "
        f"weight, {mean_crypto_weight.loc['combined_max_sharpe']:.1%} for maximum Sharpe, "
        f"{mean_crypto_weight.loc['combined_risk_parity']:.1%} for risk parity and "
        f"{mean_crypto_weight.loc['combined_min_variance']:.1%} for minimum variance. "
        "The low-volatility objective therefore rejected crypto rather than forcing "
        "diversification. Figure A4 compares risk-adjusted results: combined risk parity "
        f"led its family at {metric_by_fund.loc['combined_risk_parity','sharpe']:.2f}, "
        f"while combined maximum Sharpe was lowest at "
        f"{metric_by_fund.loc['combined_max_sharpe','sharpe']:.2f}."
    )
    selected = metrics.loc[metrics["fund_id"].isin([
        "equity_equal_weight", "crypto_risk_parity", "combined_risk_parity",
        "combined_min_variance",
    ])]
    rows = []
    for row in selected.itertuples():
        rows.append([
            row.fund_id.replace("_", " ").title(), pct(row.annualised_return),
            pct(row.annualised_volatility), f"{row.sharpe:.2f}", pct(row.max_drawdown),
        ])
    add_table(doc, ["Illustrative fund", "Return", "Vol", "Sharpe", "Max DD"], rows,
              [2.55, 0.95, 0.85, 0.75, 1.40])
    add_caption(doc, "Table 2. Selected OOS results, 2021-2023; annualised, risk-free rate 0.")
    doc.add_paragraph(
        "A fund-selection implication follows. A moderate-risk user should not choose "
        "only from return rankings: equity equal weight and combined risk parity offer "
        "stronger balance, while crypto allocations require explicit drawdown tolerance. "
        "No result is a forecast, and three OOS years are not enough to establish a "
        "permanent ranking."
    )

    doc.add_heading("4. News sentiment and structured-unstructured fusion", level=1)
    doc.add_paragraph(
        "The sentiment pipeline retains raw casing and punctuation because VADER uses "
        "intensifiers, capitals and punctuation [6]. Headlines are scored individually, "
        "aggregated to ticker-day, and then equal-weighted across tickers within each "
        "sector. Missing-news observations remain distinguishable from neutral news. "
        "The signal is lagged by one equity trading day before it can affect a weight; "
        "weekend headlines mapped to Monday are first usable on Tuesday."
    )
    doc.add_paragraph(
        "The finance lexicon extension adjusts a pre-specified set of domain words while "
        "preserving the VADER rules. It lowered the neutral classification rate by 1.76 "
        "percentage points across 146,830 headlines. That is coverage evidence, not an "
        "accuracy claim. All 24 AI-proposed terms, valences and rationales were reviewed "
        f"by the student; {int(lexicon_review['final_decision'].eq('Accept').sum())} were "
        "accepted. Financial dictionaries are useful because ordinary-language polarity "
        "can differ from financial usage [7], but human validation remains necessary."
    )
    validation = manual_validation.set_index("model")
    label_counts = manual_review["student_label"].str.lower().value_counts()
    doc.add_paragraph(
        "The stratified human review covered 50 headlines: "
        f"{int(label_counts.get('negative', 0))} negative, "
        f"{int(label_counts.get('neutral', 0))} neutral and "
        f"{int(label_counts.get('positive', 0))} positive, with mean confidence "
        f"{validation.loc['vader_enhanced','mean_student_confidence']:.2f}/5. "
        "Enhanced VADER raised overall agreement from "
        f"{validation.loc['vader_plain','agreement_rate']:.0%} to "
        f"{validation.loc['vader_enhanced','agreement_rate']:.0%}, driven by positive "
        f"agreement rising from {validation.loc['vader_plain','positive_agreement']:.1%} "
        f"to {validation.loc['vader_enhanced','positive_agreement']:.1%}. Negative "
        f"agreement remained only {validation.loc['vader_enhanced','negative_agreement']:.1%}; "
        "the extension is therefore a modest improvement, not a solved classifier."
    )
    add_table(
        doc,
        ["Model", "Overall", "Negative", "Neutral", "Positive"],
        [
            [
                model.replace("vader_", "VADER ").title(),
                pct(row.agreement_rate), pct(row.negative_agreement),
                pct(row.neutral_agreement), pct(row.positive_agreement),
            ]
            for model, row in validation.iterrows()
        ],
        [1.70, 1.20, 1.20, 1.20, 1.20],
    )
    add_caption(doc, "Table 3. Agreement with the completed 50-headline human review.")
    highest_sentiment_sector = sector_summary["mean_sentiment"].idxmax()
    lowest_sentiment_sector = sector_summary["mean_sentiment"].idxmin()
    highest_coverage_sector = sector_summary["mean_headlines"].idxmax()
    lowest_coverage_sector = sector_summary["mean_headlines"].idxmin()
    doc.add_paragraph(
        f"Figure A5 shows the sector index rather than a prediction of returns. Average "
        f"enhanced sentiment was highest for {highest_sentiment_sector} "
        f"({sector_summary.loc[highest_sentiment_sector,'mean_sentiment']:.3f}) and lowest "
        f"for {lowest_sentiment_sector} "
        f"({sector_summary.loc[lowest_sentiment_sector,'mean_sentiment']:.3f}). Coverage "
        f"was uneven: {highest_coverage_sector} averaged "
        f"{sector_summary.loc[highest_coverage_sector,'mean_headlines']:.1f} headlines per "
        f"sector-day versus {sector_summary.loc[lowest_coverage_sector,'mean_headlines']:.1f} "
        f"for {lowest_coverage_sector}. The 20-day lines clarify persistent tone, while "
        "the faint daily series shows why single-day sentiment should not be over-read."
    )
    base = fusion.set_index("variant")
    doc.add_paragraph(
        f"Figure A6 evaluates a sector-level fusion starting from the equity minimum-"
        f"variance fund: every ticker inherits its sector's lagged rolling z-score. "
        f"Fixed lambda values of 0, "
        f"+1 and -1 preserve the base, momentum and contrarian directions without tuning "
        f"on the full OOS period. The base Sharpe was {base.loc['base','sharpe']:.2f}; "
        f"momentum reached {base.loc['momentum','sharpe']:.2f}; and contrarian reached "
        f"{base.loc['contrarian','sharpe']:.2f}. Contrarian also increased annualised "
        f"return from {pct(base.loc['base','annualised_return'])} to "
        f"{pct(base.loc['contrarian','annualised_return'])} and reduced maximum drawdown "
        f"from {pct(base.loc['base','max_drawdown'])} to "
        f"{pct(base.loc['contrarian','max_drawdown'])}. Both signs are retained; the "
        "better realised sign was not retroactively selected as the model."
    )
    doc.add_paragraph(
        "The fusion result is promising but fragile. It covers one base portfolio, one "
        "market episode and an automatically proposed lexicon. Headlines may reflect "
        "events already incorporated in prices, publisher selection, and sector coverage "
        "differences. Tetlock's evidence that media tone can relate to market activity "
        "motivates testing [8], but does not establish causality here."
    )
    doc.add_page_break()

    doc.add_heading("5. Innovation and robustness", level=1)
    doc.add_heading("5.1 Ledoit-Wolf covariance shrinkage", level=2)
    doc.add_paragraph(
        "The first robustness extension replaces sample covariance only for the "
        "covariance-sensitive minimum-variance and maximum-Sharpe funds. Ledoit-Wolf "
        "shrinks the sample matrix toward a structured target using an intensity "
        "estimated within each expanding window [5]. It is not tuned against OOS "
        "performance, and the canonical 12-fund files remain byte-for-byte unchanged."
    )
    lw = shrinkage.loc[shrinkage["covariance_estimator"].eq("ledoit_wolf")]
    rows = []
    for row in lw.itertuples():
        rows.append([
            row.asset_family.title(), row.method.replace("_", " ").title(),
            f"{row.sharpe:.2f}", f"{row.delta_sharpe:+.3f}",
        ])
    add_table(doc, ["Family", "Method", "LW Sharpe", "Change"], rows,
              [1.30, 2.35, 1.35, 1.50])
    add_caption(doc, "Table 4. Ledoit-Wolf versus sample-covariance OOS Sharpe, 2021-2023.")
    doc.add_paragraph(
        "Figure A7 shows that shrinkage materially helped crypto minimum variance "
        f"({lw['delta_sharpe'].max():+.3f} Sharpe) but reduced combined minimum variance "
        f"({lw['delta_sharpe'].min():+.3f}). Its effect was smaller for the other four "
        "comparisons. The honest conclusion is conditional: covariance "
        "regularisation appears most valuable in the volatile crypto universe, but it "
        "does not uniformly dominate the baseline."
    )

    doc.add_heading("5.2 Transaction-cost sensitivity", level=2)
    doc.add_paragraph(
        "The second check deducts turnover multiplied by fixed one-way costs of 0, 10, "
        "25, 50 and 100 basis points on post-launch rebalance dates. Costs do not alter "
        "portfolio formation. The engine first lets holdings drift with returns, then "
        "measures pre-trade-to-target turnover at each monthly rebalance. At 50 bps, "
        "Figure A8 shows the largest Sharpe reduction for equity maximum Sharpe "
        f"({costs.loc[(costs['fund_id'].eq('equity_max_sharpe')) & costs['cost_bps'].eq(50),'sharpe_change_vs_gross'].iat[0]:+.3f}), "
        f"whose total turnover was {costs.loc[(costs['fund_id'].eq('equity_max_sharpe')) & costs['cost_bps'].eq(50),'total_rebalance_turnover'].iat[0]:.2f} "
        "times capital. Equal-weight funds now incur non-zero trading because relative "
        "asset returns move pre-trade holdings away from equal targets; this is more "
        "realistic than the earlier constant-target convention."
    )
    doc.add_paragraph(
        "This strengthens product interpretation. A theoretically appealing optimiser "
        "can be operationally inferior if unstable estimates drive trading. The cost "
        "model is still simplified: it omits bid-ask variation, nonlinear market impact, "
        "tax and asset-specific liquidity. It should therefore be read as sensitivity, "
        "not a precise live implementation forecast."
    )

    doc.add_heading("6. App and investor journey", level=1)
    doc.add_paragraph(
        "The six-page Streamlit journey is Home, Fund Comparison, Fund Fact Sheet, "
        "Allocation Builder, Sentiment Analytics and Robustness Lab. Home establishes "
        "scope and risk disclosure. Comparison filters the 12 funds. Fact Sheet combines "
        "return, drawdown and current target holdings. Allocation Builder enforces a 100% "
        "total and combines only shared saved OOS dates. Sentiment Analytics exposes the "
        "sector index and base-versus-tilt evidence. Robustness Lab makes estimator and "
        "cost assumptions inspectable rather than hiding them in a technical appendix."
    )
    doc.add_paragraph(
        "The application architecture is deliberately thin: it imports neither NLTK nor an "
        "optimiser and never downloads the raw data. This separates research from "
        "delivery, shortens load time, and allows app acceptance tests to confirm that "
        "every page opens from committed artifacts. The code is in the public GitHub "
        "repository https://github.com/RUOYUNWU-ui/z5652591_projectB on the main branch. "
        "The live Streamlit URL must still be verified and submitted separately."
    )

    doc.add_heading("7. Critical reflection and recommendations", level=1)
    doc.add_paragraph(
        "First, preserve simple benchmarks. Equal weight led the equity Sharpe ranking, "
        "and risk parity was competitive with lower turnover than maximum Sharpe. A live "
        "product should present optimisation as a choice with estimation risk, not an "
        "automatic upgrade. Second, treat crypto as a separately disclosed risk budget. "
        "Its drawdowns exceeded 75% even for minimum variance; user allocations should "
        "include suitability warnings and scenario-based loss communication. Third, "
        "treat the text model as a noisy contextual indicator. The completed review finds "
        "a four-percentage-point gain in overall agreement but weak negative agreement, "
        "so production use would require a larger independently labelled sample and "
        "sector-specific error analysis."
    )
    doc.add_paragraph(
        "Three further recommendations follow. (1) Extend the OOS sample before changing "
        "parameters, because three years can be dominated by one regime. (2) Add an "
        "asset-specific execution model and turnover-aware optimisation only after the "
        "current fixed-bps curves are documented. (3) Run a true untouched holdout for "
        "the sentiment tilt; the current fixed-sign comparison is transparent but not "
        "strong enough to market as predictive. I would not recommend tuning lambda, "
        "changing windows, or selecting an estimator merely to improve the displayed "
        "Sharpe. Those actions would trade honest evidence for backtest fit."
    )
    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph(
        "SignalYield meets the core product objective with 12 reproducible walk-forward "
        "funds, transparent fact-sheet evidence, lagged sector sentiment, structured-"
        "unstructured fusion and a tested investor interface. Its strongest contribution "
        "is not one winning fund but a traceable decision system: baseline assumptions "
        "remain frozen, extensions are separately labelled, negative or mixed robustness "
        "results are retained, and limitations are visible to the user."
    )

    doc.add_page_break()
    doc.add_heading("Appendix A. Required and extension exhibits", level=1)
    add_figure(doc, "performance_metrics_table.png",
               "Table A1. OOS performance metrics for all 12 funds, 2021-2023; annualised, rf=0.")
    doc.add_page_break()
    add_figure(doc, "growth_of_1_comparison.png",
               "Figure A1. Growth of $1 across funds and methods, OOS 2021-2023.")
    add_figure(doc, "drawdown_min_variance.png",
               "Figure A2. Minimum-variance drawdown by asset family, OOS 2021-2023.", 5.6)
    doc.add_page_break()
    add_figure(doc, "combined_weights_over_time.png",
               "Figure A3. Combined-fund sector and crypto weights across 36 monthly rebalances.")
    add_figure(doc, "sharpe_by_fund.png",
               "Figure A4. Annualised Sharpe ratios across the 12 funds, OOS 2021-2023.", 5.6)
    doc.add_page_break()
    add_figure(doc, "sector_sentiment_timeseries.png",
               "Figure A5. Finance-enhanced VADER sector sentiment, daily and 20-day mean, 2020-2023.")
    add_figure(doc, "fusion_before_after.png",
               "Figure A6. Equity minimum-variance base versus fixed sentiment tilts, OOS 2021-2023.", 5.5)
    doc.add_page_break()
    add_figure(doc, "covariance_shrinkage_comparison.png",
               "Figure A7. OOS Sharpe change under untuned Ledoit-Wolf covariance shrinkage.")
    add_figure(doc, "transaction_cost_sensitivity.png",
               "Figure A8. OOS Sharpe sensitivity to one-way costs from 0 to 100 bps.", 5.7)

    doc.add_page_break()
    doc.add_heading("Appendix B. Reproducibility and AI transparency", level=1)
    doc.add_paragraph(
        "The complete evidence is rebuilt by running python scripts/run_part_b.py from "
        "the repository root. It writes app data, report tables and figures. Automated "
        "tests cover portfolio timing, sentiment lagging, fusion invariants, robustness "
        "and all six app pages. scripts/check_handin.py audits required files and frozen "
        "data access. Detailed prompts, corrections and AI limitations are recorded in "
        "ai/prompt_log_part_b.md and report/QA_REPORT.md."
    )
    add_table(doc, ["Evidence", "Location"], [
        ["12-fund returns and weights", "results/data/fund_returns.csv; fund_weights.csv"],
        ["Sector sentiment", "results/data/sector_sentiment_index.csv"],
        ["Fund metrics", "results/tables/performance_metrics.csv"],
        ["Robustness tables", "results/tables/shrinkage_comparison.csv; transaction_cost_sensitivity.csv"],
        ["Completed manual review", "sentiment_manual_review_template.csv; sentiment_manual_review_validation.csv"],
    ], [2.2, 4.3])
    add_caption(doc, "Table B1. Key reproducible evidence paths.")
    doc.add_paragraph(
        "AI produced code and this editable draft under student direction. The student "
        "remains responsible for inspecting the code, checking all outputs, and rewriting "
        "the final analysis and economic interpretation in their own voice. The 50 "
        "headline labels and 24 lexicon decisions were supplied by the student and are "
        "stored separately from model predictions."
    )

    doc.add_page_break()
    doc.add_heading("References", level=1)
    references = [
        "[1] Markowitz, H. (1952). Portfolio Selection. Journal of Finance, 7(1), 77-91. https://doi.org/10.1111/j.1540-6261.1952.tb01525.x",
        "[2] Sharpe, W. F. (1966). Mutual Fund Performance. Journal of Business, 39(1), 119-138. https://doi.org/10.1086/294846",
        "[3] DeMiguel, V., Garlappi, L., and Uppal, R. (2009). Optimal Versus Naive Diversification: How Inefficient Is the 1/N Portfolio Strategy? Review of Financial Studies, 22(5), 1915-1953. https://doi.org/10.1093/rfs/hhm075",
        "[4] Maillard, S., Roncalli, T., and Teiletche, J. (2010). The Properties of Equally Weighted Risk Contribution Portfolios. Journal of Portfolio Management, 36(4), 60-70. https://doi.org/10.3905/jpm.2010.36.4.060",
        "[5] Ledoit, O., and Wolf, M. (2004). A Well-Conditioned Estimator for Large-Dimensional Covariance Matrices. Journal of Multivariate Analysis, 88(2), 365-411. https://doi.org/10.1016/S0047-259X(03)00096-4",
        "[6] Hutto, C. J., and Gilbert, E. (2014). VADER: A Parsimonious Rule-Based Model for Sentiment Analysis of Social Media Text. Proceedings of ICWSM, 8(1), 216-225. https://doi.org/10.1609/icwsm.v8i1.14550",
        "[7] Loughran, T., and McDonald, B. (2011). When Is a Liability Not a Liability? Textual Analysis, Dictionaries, and 10-Ks. Journal of Finance, 66(1), 35-65. https://doi.org/10.1111/j.1540-6261.2010.01625.x",
        "[8] Tetlock, P. C. (2007). Giving Content to Investor Sentiment: The Role of Media in the Stock Market. Journal of Finance, 62(3), 1139-1168. https://doi.org/10.1111/j.1540-6261.2007.01232.x",
        "[9] Bailey, D. H., Borwein, J. M., Lopez de Prado, M., and Zhu, Q. J. (2014). The Probability of Backtest Overfitting. Journal of Computational Finance, 20(4), 39-69. https://doi.org/10.21314/JCF.2016.322",
    ]
    for reference in references:
        paragraph = doc.add_paragraph(reference)
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.space_after = Pt(5)
        for run in paragraph.runs:
            run.font.size = Pt(9)

    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"wrote {OUTPUT}")
    return OUTPUT


if __name__ == "__main__":
    build()
